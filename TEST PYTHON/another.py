from docx import Document

def sql_to_docx(sql_file_path, output_docx_path):
    document = Document()
    document.add_heading('Database Schema: job_portail', 0)

    with open(sql_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table = None
    for line in lines:
        line = line.strip()
        if line.startswith("CREATE TABLE"):
            table = line.split('`')[1]
            document.add_heading(f"🧾 Table: {table}", level=1)
        elif table and line.startswith("`"):
            parts = line.split()
            field = parts[0].strip('`')
            type_ = parts[1]
            document.add_paragraph(f"• {field}: {type_}", style='List Bullet')
        elif line.startswith(")"):
            table = None

    document.save(output_docx_path)
    print(f"✅ DOCX exported: {output_docx_path}")

# Usage
sql_to_docx("job_portail_schema.sql", "job_portail_schema.docx")
